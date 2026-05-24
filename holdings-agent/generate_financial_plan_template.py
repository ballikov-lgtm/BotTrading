"""
generate_financial_plan_template.py  —  one-shot helper

Creates a sample FinancialPlan.docx in the private folder. The allocator
(allocator.py via weekly_review.py) reads this each weekly review run.

The plan governs HOW the agent allocates capital — your cash on hand, your
risk preference, position caps, preferred wrapper. Edit as your situation changes.

Run:
    py generate_financial_plan_template.py
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PRIVATE_DIR = Path(os.environ.get(
    "HOLDINGS_PRIVATE_DIR",
    str(Path.home() / "OneDrive" / "Documents" / "Private Investments"),
))
OUTPUT_PATH = PRIVATE_DIR / "FinancialPlan.docx"


CASH_PREFS_HEADERS  = ["Field", "Value"]
CASH_PREFS_DEFAULTS = [
    ["Cash available to invest (GBP)",  "10000"],
    ["Risk target",                     "MEDIUM"],
    ["Max single position (%)",         "25"],
    ["Preferred wrapper",               "ISA"],
    ["Minimum allocation chunk (GBP)",  "500"],
]

RULES_HEADERS  = ["Field", "Value"]
RULES_DEFAULTS = [
    ["Top-up existing INSUFFICIENT_DATA holdings",   "Only passive / index funds"],
    ["Maximum number of new positions per review",   "6"],
    ["Watchlist bonus weighting",                    "Slight bias toward watchlist matches"],
    ["Notes",                                        "Free text the allocator passes into its reasoning prompt"],
]


def add_keyvalue_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    for r, row_data in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            if i == 0:
                for run in cells[i].paragraphs[0].runs:
                    run.bold = True
    return table


def main():
    PRIVATE_DIR.mkdir(parents=True, exist_ok=True)

    if OUTPUT_PATH.exists():
        print(f"REFUSING to overwrite existing file: {OUTPUT_PATH}")
        print("Delete or rename it first if you want a fresh template.")
        return 1

    doc = Document()

    title = doc.add_heading("My Financial Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "This is your private allocation policy. The agent's weekly review reads "
        "this file each run to compute a fresh allocation recommendation for any "
        "available capital (cash + SELL proceeds). Edit when your situation changes "
        "— e.g. after you deposit fresh cash, change risk appetite, or move to a "
        "different wrapper. Never share this document."
    ).italic = True

    doc.add_heading("Cash & Preferences", level=1)
    doc.add_paragraph(
        "The core inputs to the allocation. \"Cash available\" should be the "
        "amount sitting uninvested today (excluding emergency reserves). \"Risk "
        "target\" steers the agent's split: LOW = ~90% defensive, LOW-MED = ~80%, "
        "MEDIUM = ~65%, MED-HIGH = ~55%, HIGH = ~50%. \"Max single position\" "
        "caps how much of any one allocation can go into a single fund."
    )
    add_keyvalue_table(doc, CASH_PREFS_HEADERS, CASH_PREFS_DEFAULTS)

    doc.add_paragraph()

    doc.add_heading("Allocation Rules", level=1)
    doc.add_paragraph(
        "Additional rules the allocator follows. The Notes field passes through "
        "into the allocator's reasoning prompt — use it for any context you want "
        "the agent to weigh (e.g. \"close to retirement, prefer income\", or "
        "\"keep at least 10% in gold-related funds\")."
    )
    add_keyvalue_table(doc, RULES_HEADERS, RULES_DEFAULTS)

    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Notes: Risk target values: LOW, LOW-MED, MEDIUM, MED-HIGH, HIGH. "
        "Wrapper values: ISA, SIPP, GIA (General Investment Account). "
        "All cash amounts in GBP. "
        "Defaults above (£10,000 cash, MEDIUM target, 25% max) are sensible starting points — "
        "update as needed and save."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT_PATH)
    print(f"Created template: {OUTPUT_PATH}")
    print("Defaults pre-filled (£10,000 cash, MEDIUM target). Edit in Word as needed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
