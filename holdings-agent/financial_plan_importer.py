"""
holdings-agent/financial_plan_importer.py — reads FinancialPlan.docx into a dict.

Two key-value tables: Cash & Preferences, Allocation Rules.

Returns None if file missing (caller treats as "no plan on file — skip allocator").
"""

import os
import re
from pathlib import Path

from docx import Document


VALID_RISK = {"LOW", "LOW-MED", "MEDIUM", "MED-HIGH", "HIGH"}
VALID_WRAPPERS = {"ISA", "SIPP", "GIA"}


def _strip_number(s: str):
    if not s:
        return None
    cleaned = re.sub(r'[^\d.\-]', '', s)
    if not cleaned or cleaned in ("-", "."):
        return None
    try:
        return float(cleaned)
    except ValueError:
        return None


def _table_to_dict(table) -> dict:
    out = {}
    rows = list(table.rows)
    if not rows:
        return out
    start = 1 if rows[0].cells[0].text.strip().lower() == "field" else 0
    for row in rows[start:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            out[cells[0]] = cells[1]
    return out


def import_plan_from_docx(docx_path: Path):
    """Returns plan dict or None if file doesn't exist."""
    if not docx_path.exists():
        return None

    doc = Document(str(docx_path))
    if len(doc.tables) < 2:
        raise ValueError(
            f"FinancialPlan.docx must contain at least 2 tables (Cash & Preferences + Allocation Rules). "
            f"Found {len(doc.tables)}."
        )

    prefs = _table_to_dict(doc.tables[0])
    rules = _table_to_dict(doc.tables[1])

    risk_raw = (prefs.get("Risk target") or "MEDIUM").strip().upper()
    if risk_raw not in VALID_RISK:
        risk_raw = "MEDIUM"

    wrapper_raw = (prefs.get("Preferred wrapper") or "ISA").strip().upper()
    if wrapper_raw not in VALID_WRAPPERS:
        wrapper_raw = "ISA"

    max_topup_rule = (rules.get("Top-up existing INSUFFICIENT_DATA holdings") or "").strip().lower()
    allow_topup_insufficient = (
        "passive" not in max_topup_rule  # default: only top up passive/index
        and "no" not in max_topup_rule
    )
    topup_passive_only = "passive" in max_topup_rule or "index" in max_topup_rule

    plan = {
        "cash_available_gbp":          _strip_number(prefs.get("Cash available to invest (GBP)", "0")) or 0.0,
        "risk_target":                 risk_raw,
        "max_single_position_pct":     _strip_number(prefs.get("Max single position (%)", "25")) or 25.0,
        "preferred_wrapper":           wrapper_raw,
        "min_allocation_chunk_gbp":    _strip_number(prefs.get("Minimum allocation chunk (GBP)", "500")) or 500.0,
        "topup_rule":                  rules.get("Top-up existing INSUFFICIENT_DATA holdings", "") or "",
        "topup_passive_only":          topup_passive_only,
        "max_new_positions":           int(_strip_number(rules.get("Maximum number of new positions per review", "6")) or 6),
        "watchlist_bonus_rule":        rules.get("Watchlist bonus weighting", "") or "",
        "user_notes":                  rules.get("Notes", "") or "",
        "_source":                     docx_path.name,
    }

    # If cash is 0 AND user hasn't customised the placeholder-y notes, flag as unfilled
    plan["_looks_unfilled"] = (
        plan["cash_available_gbp"] == 0
        and "free text" in plan["user_notes"].lower()
    )

    return plan


if __name__ == "__main__":
    import json
    private_dir = Path(os.environ.get(
        "HOLDINGS_PRIVATE_DIR",
        str(Path.home() / "OneDrive" / "Documents" / "Private Investments"),
    ))
    result = import_plan_from_docx(private_dir / "FinancialPlan.docx")
    if result is None:
        print(f"No FinancialPlan.docx found.")
    else:
        print(json.dumps(result, indent=2, ensure_ascii=False))
