"""
generate_pension_template.py  —  one-shot helper

Creates a sample Pension.docx in the private folder. The pension importer
(pension_importer.py) reads this file each weekly review run.

This file is NOT part of the agent runtime — it's a setup utility.

Run:
    py generate_pension_template.py
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PRIVATE_DIR = Path(os.environ.get(
    "HOLDINGS_PRIVATE_DIR",
    str(Path.home() / "OneDrive" / "Documents" / "Private Investments"),
))
OUTPUT_PATH = PRIVATE_DIR / "Pension.docx"


PENSION_DETAILS_HEADERS  = ["Field", "Value"]
PENSION_DETAILS_EXAMPLES = [
    ["Provider",                         "e.g. Aviva, Standard Life, Scottish Widows"],
    ["Plan name",                        "e.g. Workplace pension, SIPP"],
    ["Current value (GBP)",              "0"],
    ["Expected annual return (%)",       "4.0"],
    ["Monthly contribution (GBP)",       "0"],
    ["Employer monthly contribution (GBP)", "0"],
]

RETIREMENT_HEADERS  = ["Field", "Value"]
RETIREMENT_EXAMPLES = [
    ["Date of birth",                    "YYYY-MM-DD"],
    ["Target retirement age",            "62"],
    ["Target retirement pot (GBP)",      "0"],
    ["Target annual income in retirement (GBP)", "0"],
]


def add_keyvalue_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    for r, row_data in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            if i == 0:  # field name column bold
                for run in cells[i].paragraphs[0].runs:
                    run.bold = True
    return table


def main():
    PRIVATE_DIR.mkdir(parents=True, exist_ok=True)

    if OUTPUT_PATH.exists():
        print(f"REFUSING to overwrite existing file: {OUTPUT_PATH}")
        print("Delete or rename it first if you want a fresh template.")
        return 1

    doc = Document()

    title = doc.add_heading("My Pension & Retirement Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "This is your private pension and retirement record. The weekly review reads "
        "this file to project your trajectory toward retirement, combined with your "
        "actively-managed holdings (Holdings.docx). Edit the tables below — replace the "
        "example values with your real numbers. Save when done. Never share this document."
    ).italic = True

    doc.add_heading("Pension Details", level=1)
    doc.add_paragraph(
        "One row per field. The pension is treated as a single slow-growth holding (you can't "
        "sell-and-rebuy individual funds inside it), so no SELL/HOLD/BUY recommendations are "
        "made — but the value is included in total wealth and the retirement projection."
    )
    add_keyvalue_table(doc, PENSION_DETAILS_HEADERS, PENSION_DETAILS_EXAMPLES)

    doc.add_paragraph()

    doc.add_heading("Retirement Planning", level=1)
    doc.add_paragraph(
        "Used to project where your combined portfolio (holdings + pension) will be at "
        "your target retirement age, vs your target pot. If you're behind, the weekly "
        "review will calculate the extra monthly contribution needed to close the gap."
    )
    add_keyvalue_table(doc, RETIREMENT_HEADERS, RETIREMENT_EXAMPLES)

    doc.add_paragraph()

    doc.add_heading("Notes", level=1)
    notes = doc.add_paragraph(
        "Free text — anything else relevant about your pension, retirement plans, "
        "expected state pension, planned downsize, inheritance considerations, etc. "
        "This text passes through into the weekly review as context but is not parsed."
    )

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Notes: dates in ISO format (YYYY-MM-DD). All monetary values in GBP. "
        "Expected annual return is a single percentage (e.g. 4.0 for 4%). "
        "Do not add or remove rows in the tables — the importer expects the exact field names above."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT_PATH)
    print(f"Created template: {OUTPUT_PATH}")
    print("Open it in Word, replace the example values with your real numbers, save.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
