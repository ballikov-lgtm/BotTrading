"""
holdings-agent/pension_importer.py — reads Pension.docx into pension.json.

Pension.docx layout (two key-value tables + a free-text Notes section):

  Table 1 — Pension Details
    Provider                          | <text>
    Plan name                         | <text>
    Current value (GBP)               | <number>
    Expected annual return (%)        | <number, e.g. 4.0>
    Monthly contribution (GBP)        | <number>
    Employer monthly contribution (GBP)| <number>

  Table 2 — Retirement Planning
    Date of birth                     | YYYY-MM-DD
    Target retirement age             | <integer>
    Target retirement pot (GBP)       | <number>
    Target annual income in retirement (GBP) | <number>

Importer is forgiving:
  - Missing optional fields default to None.
  - Currency / "%" / commas in numbers are stripped.
  - If Pension.docx isn't present, import_pension_from_docx returns None (caller
    treats it as "no pension data on file").
"""

import os
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document


def _strip_number(s: str):
    """Pull a numeric value out of a string. '£12,000' -> 12000.0, '4%' -> 4.0."""
    if not s:
        return None
    cleaned = re.sub(r'[^\d.\-]', '', s)
    if not cleaned or cleaned in ("-", "."):
        return None
    try:
        return float(cleaned)
    except ValueError:
        return None


def _parse_date(s: str):
    """Accept YYYY-MM-DD only (the template tells the user this)."""
    if not s:
        return None
    try:
        return date.fromisoformat(s.strip())
    except ValueError:
        return None


def _table_to_dict(table) -> dict:
    """Two-column key-value table → dict, with raw string values."""
    out = {}
    rows = list(table.rows)
    if not rows:
        return out
    # Skip header row if it looks like a header (first cell == "Field")
    start = 1 if rows[0].cells[0].text.strip().lower() == "field" else 0
    for row in rows[start:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            out[cells[0]] = cells[1]
    return out


def _years_until(target_date: date, today_d: date = None) -> float:
    today_d = today_d or date.today()
    delta_days = (target_date - today_d).days
    return delta_days / 365.25


def import_pension_from_docx(docx_path: Path):
    """Returns pension dict or None if file doesn't exist."""
    if not docx_path.exists():
        return None

    doc = Document(str(docx_path))
    if len(doc.tables) < 2:
        raise ValueError(
            f"Pension.docx must contain at least 2 tables (Pension Details + Retirement Planning). "
            f"Found {len(doc.tables)}."
        )

    details   = _table_to_dict(doc.tables[0])
    planning  = _table_to_dict(doc.tables[1])

    dob = _parse_date(planning.get("Date of birth", ""))

    pension = {
        "schema_version":               "0.1",
        "last_updated":                 date.today().isoformat(),
        "provider":                     details.get("Provider") or None,
        "plan_name":                    details.get("Plan name") or None,
        "current_value_gbp":            _strip_number(details.get("Current value (GBP)", "")),
        "expected_annual_return_pct":   _strip_number(details.get("Expected annual return (%)", "")),
        "monthly_contribution_gbp":     _strip_number(details.get("Monthly contribution (GBP)", "")),
        "employer_monthly_contribution_gbp": _strip_number(details.get("Employer monthly contribution (GBP)", "")),
        "date_of_birth":                dob.isoformat() if dob else None,
        "target_retirement_age":        int(_strip_number(planning.get("Target retirement age", "")) or 0) or None,
        "target_retirement_pot_gbp":    _strip_number(planning.get("Target retirement pot (GBP)", "")),
        "target_annual_income_gbp":     _strip_number(planning.get("Target annual income in retirement (GBP)", "")),
        "_source":                      docx_path.name,
    }

    # Validation flags so the caller can show "Pension.docx still has placeholder values"
    placeholder_strings = {"e.g. aviva, standard life, scottish widows", "e.g. workplace pension, sipp", "yyyy-mm-dd"}
    pension["_looks_unfilled"] = (
        (pension["current_value_gbp"] in (None, 0))
        and (pension["target_retirement_pot_gbp"] in (None, 0))
        and (not pension["date_of_birth"])
    )

    return pension


if __name__ == "__main__":
    # Standalone test: read the user's Pension.docx and print the parsed result
    import json
    private_dir = Path(os.environ.get(
        "HOLDINGS_PRIVATE_DIR",
        str(Path.home() / "OneDrive" / "Documents" / "Private Investments"),
    ))
    path = private_dir / "Pension.docx"
    result = import_pension_from_docx(path)
    if result is None:
        print(f"No Pension.docx at {path}")
    else:
        print(json.dumps(result, indent=2, ensure_ascii=False))
