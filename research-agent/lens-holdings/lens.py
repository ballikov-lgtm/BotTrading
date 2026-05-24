"""
lens-holdings/lens.py  —  Holdings research lens (v0.2)

PRIVACY: real holdings + generated reports live in a PRIVATE folder OUTSIDE this
git repo (default: ~/OneDrive/Documents/Private Investments/). The repo only
contains the lens code + the (public) sources list. The user's holdings .docx,
the derived holdings.json, and all output reports never enter the working tree.

Flow:
  1. If Holdings.docx exists in the private folder, import it → holdings.json
  2. Use Claude with web_search to build shared context (macro + geopolitical + 7-day calendar)
  3. Research each holding individually with a GREEN/AMBER/RED signal
  4. Scan for new investment opportunities (named funds / managers)
  5. Write holdings-alerts.json into the private reports folder

No email — that's Phase 4.

Override the private folder location with HOLDINGS_PRIVATE_DIR in .env if needed.

Run:
    pip install -r requirements.txt
    python lens.py
"""

import json
import os
import re
import shutil
import subprocess
import sys
import time
import traceback
from datetime import date, datetime, timezone
from pathlib import Path

from anthropic import Anthropic
from docx import Document
from dotenv import load_dotenv

# ─── PATHS ────────────────────────────────────────────────────────────────────

LENS_DIR             = Path(__file__).parent
RESEARCH_AGENT_DIR   = LENS_DIR.parent
TRADING_SETUP_DIR    = RESEARCH_AGENT_DIR.parent
ENV_FILE             = TRADING_SETUP_DIR / ".env"

# Load .env BEFORE resolving PRIVATE_DIR so an override there is honoured.
load_dotenv(ENV_FILE, override=True)  # override=True: shell may have ANTHROPIC_API_KEY=""

PRIVATE_DIR = Path(os.environ.get(
    "HOLDINGS_PRIVATE_DIR",
    str(Path.home() / "OneDrive" / "Documents" / "Private Investments"),
))
HOLDINGS_DOCX  = PRIVATE_DIR / "Holdings.docx"
HOLDINGS_FILE  = PRIVATE_DIR / "holdings.json"
OUTPUT_FILE    = PRIVATE_DIR / "reports" / "holdings-alerts.json"

SOURCES_FILE   = LENS_DIR / "sources.json"  # public — research domains, no PII

# ─── CONFIG ───────────────────────────────────────────────────────────────────

MODEL      = "claude-sonnet-4-6"     # research-quality, lower cost than Opus
GENERATOR  = "lens-holdings v0.4.0"

# User's target annualised return range — used to classify holdings vs target.
# If you change these, the holding research prompt picks them up automatically.
TARGET_ANNUAL_RETURN_MIN = 10.0   # %  (UNDERPERFORMING if 10yr < this)
TARGET_ANNUAL_RETURN_MAX = 14.0   # %  (OUTPERFORMING  if 10yr > this)
WEB_TOOL   = {"type": "web_search_20250305", "name": "web_search"}

# Web-search uses per call. Each search result becomes input tokens on the next
# model turn, so higher values consume the 30K TPM tier limit fast.
USES_SHARED       = 3
USES_PER_HOLDING  = 2
USES_OPPORTUNITY  = 3

# Pause between API calls to stay under the 30K input tokens-per-minute tier limit.
# Tune down if/when the account is on a higher rate tier — current entry tier needs ~60s
# to fully clear after a heavy web_search call.
PAUSE_SEC = 60

client = Anthropic()

# ─── HELPERS ──────────────────────────────────────────────────────────────────

def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def collect_text(response) -> str:
    """Join all text-content blocks from a Messages API response."""
    return "".join(
        block.text for block in response.content if hasattr(block, "text")
    )


def extract_json(text: str) -> dict:
    """Pull the first JSON object out of a text blob.

    Robust to: ```json fenced blocks, surrounding prose, prose AFTER the JSON
    that happens to contain braces. Walks the brace depth honoring string escapes.
    """
    fenced = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', text, re.DOTALL)
    if fenced:
        return json.loads(fenced.group(1))

    start = text.find('{')
    if start < 0:
        raise ValueError(f"No JSON object found. First 300 chars: {text[:300]}")

    depth = 0
    in_string = False
    escape = False
    for i in range(start, len(text)):
        c = text[i]
        if escape:
            escape = False
            continue
        if c == '\\':
            escape = True
            continue
        if c == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if c == '{':
            depth += 1
        elif c == '}':
            depth -= 1
            if depth == 0:
                return json.loads(text[start:i + 1])

    raise ValueError(f"Unbalanced braces from char {start}. First 300: {text[start:start + 300]}")


def domains_for_prompt(source_list: list) -> str:
    return ", ".join(s["domain"] for s in source_list)


def web_search_tool(max_uses: int) -> dict:
    return {**WEB_TOOL, "max_uses": max_uses}


# ─── HOLDINGS.DOCX IMPORTER ──────────────────────────────────────────────────

def slugify(name: str) -> str:
    s = re.sub(r'[^a-z0-9]+', '-', name.lower()).strip('-')
    return s[:60] or "unnamed"


def _row_cells(row) -> list:
    return [cell.text.strip() for cell in row.cells]


def _parse_holdings_table(table) -> list:
    rows = list(table.rows)
    if not rows:
        return []
    headers = _row_cells(rows[0])
    h_idx = {h: i for i, h in enumerate(headers)}

    required = ["Name", "Type", "Value GBP", "Purchased", "Standing Alerts", "Notes"]
    missing = [h for h in required if h not in h_idx]
    if missing:
        raise ValueError(f"Holdings table missing columns: {missing}. Found: {headers}")

    holdings = []
    for row in rows[1:]:
        cells = _row_cells(row)
        name = cells[h_idx["Name"]]
        if not name:
            continue  # skip blank rows

        value_str = cells[h_idx["Value GBP"]].replace(",", "")
        try:
            value_gbp = int(float(value_str)) if value_str else None
        except ValueError:
            value_gbp = None

        alerts_str = cells[h_idx["Standing Alerts"]]
        alerts = [a.strip() for a in alerts_str.split(",") if a.strip()] if alerts_str else []

        holdings.append({
            "id":              slugify(name),
            "name":            name,
            "type":            cells[h_idx["Type"]] or "Other",
            "value_gbp":       value_gbp,
            "purchased_at":    cells[h_idx["Purchased"]] or None,
            "standing_alerts": alerts,
            "notes":           cells[h_idx["Notes"]] or "",
        })
    return holdings


def _parse_watchlist_table(table) -> list:
    rows = list(table.rows)
    if not rows:
        return []
    headers = _row_cells(rows[0])
    h_idx = {h: i for i, h in enumerate(headers)}

    required = ["Name", "Type", "Rationale"]
    missing = [h for h in required if h not in h_idx]
    if missing:
        raise ValueError(f"Watchlist table missing columns: {missing}. Found: {headers}")

    watchlist = []
    for row in rows[1:]:
        cells = _row_cells(row)
        name = cells[h_idx["Name"]]
        if not name:
            continue
        watchlist.append({
            "id":        slugify(name),
            "name":      name,
            "type":      cells[h_idx["Type"]] or "Other",
            "rationale": cells[h_idx["Rationale"]] or "",
        })
    return watchlist


def import_holdings_from_docx(docx_path: Path) -> dict:
    """Read the two tables (holdings + watchlist) out of Holdings.docx."""
    doc = Document(str(docx_path))
    if len(doc.tables) < 2:
        raise ValueError(
            f"Expected at least 2 tables in {docx_path.name}, found {len(doc.tables)}. "
            "Make sure the file has a Current Holdings table AND a Watchlist table."
        )
    return {
        "schema_version": "0.1",
        "last_updated":   date.today().isoformat(),
        "currency":       "GBP",
        "holdings":       _parse_holdings_table(doc.tables[0]),
        "watchlist":      _parse_watchlist_table(doc.tables[1]),
        "_source":        f"imported from {docx_path.name}",
    }


# ─── REPORT ARCHIVE + GIT SYNC ───────────────────────────────────────────────

def archive_report(output_file: Path) -> Path:
    """Copy holdings-alerts.json to archive/holdings-alerts-YYYY-MM-DD.json."""
    archive_dir = output_file.parent / "archive"
    archive_dir.mkdir(parents=True, exist_ok=True)
    archive_path = archive_dir / f"holdings-alerts-{date.today().isoformat()}.json"
    shutil.copy2(output_file, archive_path)
    return archive_path


def push_report_to_github(reports_dir: Path) -> None:
    """If reports_dir is a git repo, commit + push. Soft-fails on any error so a
    network glitch or auth issue never breaks the lens run itself."""
    if not (reports_dir / ".git").exists():
        return  # not a git repo — nothing to do

    def git(*args, timeout=60):
        return subprocess.run(
            ["git", "-C", str(reports_dir), *args],
            capture_output=True, text=True, timeout=timeout,
        )

    try:
        add = git("add", ".")
        if add.returncode != 0:
            print(f"  ! git add failed: {add.stderr.strip()[:200]}")
            return

        # Anything actually staged?
        diff = git("diff", "--cached", "--quiet")
        if diff.returncode == 0:
            print("  [OK] no changes to push (report unchanged)")
            return

        commit_msg = f"Daily holdings report {date.today().isoformat()}"
        commit = git("commit", "-m", commit_msg)
        if commit.returncode != 0:
            print(f"  ! git commit failed: {commit.stderr.strip()[:200]}")
            return

        push = git("push", "origin", "main", timeout=120)
        if push.returncode != 0:
            print(f"  ! git push failed: {push.stderr.strip()[:200]}")
        else:
            print(f"  [OK] pushed to github.com (commit: {commit_msg})")
    except subprocess.TimeoutExpired:
        print("  ! git sync timed out")
    except Exception as exc:
        print(f"  ! git sync error: {exc}")


# ─── RESEARCH CALLS ───────────────────────────────────────────────────────────

def research_shared_context(sources: dict) -> dict:
    """Macro + geopolitical + 7-day event calendar. Universal — would feed all lenses."""
    today    = datetime.now().strftime("%A %d %B %Y")
    preferred = domains_for_prompt(sources["default_sources"])

    response = client.messages.create(
        model=MODEL,
        max_tokens=2000,
        tools=[web_search_tool(USES_SHARED)],
        system=(
            "You are a macro research analyst briefing a UK private investor holding global "
            "equity funds and ETFs. Use the web_search tool to gather LIVE, current data. "
            f"Prefer these sources: {preferred}. "
            "Return JSON only (no surrounding prose) in this exact shape: "
            '{"macro_summary": "...", '
            '"geopolitical_summary": "...", '
            '"calendar_next_7d": [{"date": "YYYY-MM-DD", "event": "...", "importance": "HIGH|MEDIUM|LOW"}]}'
        ),
        messages=[{
            "role": "user",
            "content": (
                f"Today is {today}. Provide:\n"
                "1. Macro briefing — Fed/BoE/ECB recent moves, inflation trajectory, equity regime\n"
                "2. Geopolitical briefing — wars, sanctions, major political risks affecting markets\n"
                "3. Calendar of major events in the next 7 sessions (FOMC, BoE, NFP, CPI, major earnings)\n\n"
                "Be concise but factual. Return JSON only."
            )
        }]
    )
    return extract_json(collect_text(response))


def research_holding(holding: dict, sources: dict) -> dict:
    """Per-holding live research with GREEN/AMBER/RED signal AND performance data."""
    today     = datetime.now().strftime("%A %d %B %Y")
    preferred = domains_for_prompt(sources["default_sources"])
    standing  = "; ".join(holding.get("standing_alerts", [])) or "none"
    tmin      = f"{TARGET_ANNUAL_RETURN_MIN:g}"
    tmax      = f"{TARGET_ANNUAL_RETURN_MAX:g}"

    response = client.messages.create(
        model=MODEL,
        max_tokens=1800,
        tools=[web_search_tool(USES_PER_HOLDING)],
        system=(
            "You are a fund research analyst. Use web_search to gather LIVE current data "
            "for the named fund. Capture three things:\n"
            "1. NEWS — manager changes, ratings changes, regulatory issues, sector exposure shifts\n"
            "2. PERFORMANCE — annualised total returns for 1yr, 3yr, 5yr, 10yr (sourced from "
            "Trustnet, Hargreaves Lansdown, or Morningstar — these all publish standard annualised figures)\n"
            "3. BENCHMARK — the appropriate sector/index benchmark and ITS 10yr annualised return for comparison\n\n"
            f"User's target return: {tmin}-{tmax}% annualised. Classify the fund vs target:\n"
            f"  UNDERPERFORMING — 10yr annualised < {tmin}%\n"
            f"  ON_TRACK       — 10yr annualised between {tmin}% and {tmax}%\n"
            f"  OUTPERFORMING  — 10yr annualised > {tmax}%\n"
            "  INSUFFICIENT_DATA — fund younger than 10yr (fall back to longest available period; explain in summary)\n\n"
            "Rate the SIGNAL based on news AND performance combined:\n"
            "  GREEN — no concerns, on-track or outperforming\n"
            "  AMBER — worth monitoring (mildly underperforming, or notable news, or watchlist concerns)\n"
            "  RED   — act/investigate urgently (sustained underperformance vs target, or material risk event)\n\n"
            f"Prefer these sources: {preferred}. "
            "If a performance figure isn't available, use null — do not invent.\n"
            "Return JSON only, no surrounding prose:\n"
            '{"signal": "GREEN|AMBER|RED", "summary": "2-3 sentences combining news + performance", '
            '"bullets": ["...", "..."], "sources_used": ["domain1", "domain2"], '
            '"performance": {'
            '"annualised_returns_pct": {"1yr": null, "3yr": null, "5yr": null, "10yr": null}, '
            '"benchmark": {"name": "...", "annualised_10yr_pct": null}, '
            '"vs_target": "UNDERPERFORMING|ON_TRACK|OUTPERFORMING|INSUFFICIENT_DATA", '
            f'"target_range_pct": "{tmin}-{tmax}"'
            '}}'
        ),
        messages=[{
            "role": "user",
            "content": (
                f"Today is {today}.\n"
                f"Fund: {holding['name']}\n"
                f"Type: {holding.get('type', 'unknown')}\n"
                f"Standing alerts: {standing}\n\n"
                "Gather live news AND historical performance data for this fund. Return JSON only."
            )
        }]
    )
    result = extract_json(collect_text(response))
    result["id"]   = holding["id"]
    result["name"] = holding["name"]
    return result


def research_opportunities(sources: dict) -> list:
    """Scan for funds / managers predicted to do well — the 'where should I deploy cash' question."""
    today     = datetime.now().strftime("%A %d %B %Y")
    opp_list  = sources.get("opportunity_sources", sources["default_sources"])
    preferred = domains_for_prompt(opp_list)

    response = client.messages.create(
        model=MODEL,
        max_tokens=5000,
        tools=[web_search_tool(USES_OPPORTUNITY)],
        system=(
            "You are an investment ideas analyst for a UK private investor with cash ready "
            "to deploy. Use web_search to find funds, ETFs, or fund managers predicted to "
            "do well over the next 6-12 months. Focus on: newly-rated HL Wealth Shortlist "
            "additions, Morningstar analyst picks, Citywire AAA-rated managers, top-decile "
            "performers in their sector with a credible forward outlook. "
            f"Prefer these sources: {preferred}. "
            "Avoid generic advice ('invest in the S&P 500'). Find SPECIFIC, NAMED opportunities "
            "with a clear rationale. Return JSON only: "
            '{"opportunities": [{"name": "...", "type": "OEIC|ETF|InvestmentTrust|Equity", '
            '"rationale": "2-3 sentences", "predicted_outlook": "...", '
            '"risk_level": "LOW|MEDIUM|HIGH", '
            '"bullets": ["..."], "sources_used": ["..."]}]}'
        ),
        messages=[{
            "role": "user",
            "content": (
                f"Today is {today}. Surface 3-5 specific named investment opportunities "
                "for a UK retail investor with global equity exposure already. Return JSON only."
            )
        }]
    )
    parsed = extract_json(collect_text(response))
    return parsed.get("opportunities", [])


# ─── SUMMARY ──────────────────────────────────────────────────────────────────

def summarise_alerts(holdings_results: list, opportunities: list) -> dict:
    counts = {"red_count": 0, "amber_count": 0, "green_count": 0}
    for h in holdings_results:
        sig = h.get("signal", "AMBER").upper()
        key = f"{sig.lower()}_count"
        counts[key] = counts.get(key, 0) + 1
    counts["new_opportunities_count"] = len(opportunities)
    return counts


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main() -> int:
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print(f"ERROR: ANTHROPIC_API_KEY not in environment. Looked for .env at {ENV_FILE}")
        return 1

    started_at = datetime.now()
    print(f"[{started_at.isoformat(timespec='seconds')}] lens-holdings starting...")
    print(f"  Private data folder: {PRIVATE_DIR}")

    # Step 0: import from .docx if present (it's the user's source of truth)
    if HOLDINGS_DOCX.exists():
        print(f"  Importing from {HOLDINGS_DOCX.name}...")
        try:
            imported = import_holdings_from_docx(HOLDINGS_DOCX)
            HOLDINGS_FILE.parent.mkdir(parents=True, exist_ok=True)
            with HOLDINGS_FILE.open("w", encoding="utf-8") as f:
                json.dump(imported, f, indent=2, ensure_ascii=False)
            print(f"  Wrote {HOLDINGS_FILE.name} "
                  f"({len(imported['holdings'])} holdings, {len(imported['watchlist'])} watchlist)")
        except Exception as exc:
            print(f"  ! .docx import failed: {exc}")
            print(f"  Falling back to existing {HOLDINGS_FILE.name} if present...")
    else:
        print(f"  No {HOLDINGS_DOCX.name} found — using existing {HOLDINGS_FILE.name} if present.")

    if not HOLDINGS_FILE.exists():
        print(f"ERROR: No holdings data. Either place a Holdings.docx in {PRIVATE_DIR}")
        print(f"  or create {HOLDINGS_FILE} by hand.")
        return 1

    holdings_data = load_json(HOLDINGS_FILE)
    sources_data  = load_json(SOURCES_FILE)
    holdings = holdings_data.get("holdings", [])

    print(f"  Loaded {len(holdings)} holding(s) from {HOLDINGS_FILE.name}")

    def pause(label: str):
        print(f"  ... pausing {PAUSE_SEC}s before {label} (rate-limit safety)")
        time.sleep(PAUSE_SEC)

    print("  Researching shared context (macro + geopolitical + calendar)...")
    try:
        shared = research_shared_context(sources_data)
    except Exception as exc:
        print(f"  ! shared context failed: {exc}")
        shared = {"error": str(exc)}

    print(f"  Researching {len(holdings)} holding(s) individually...")
    holdings_results = []
    for i, h in enumerate(holdings):
        if i == 0:
            pause("first holding")
        else:
            pause(f"holding {i + 1}")
        print(f"    - {h['name']}")
        try:
            holdings_results.append(research_holding(h, sources_data))
        except Exception as exc:
            print(f"      ! failed: {exc}")
            holdings_results.append({
                "id":      h["id"],
                "name":    h["name"],
                "signal":  "AMBER",
                "summary": f"Research failed: {exc}",
                "bullets": [],
                "sources_used": [],
            })

    pause("opportunity scan")
    print("  Scanning for new opportunities...")
    try:
        opportunities = research_opportunities(sources_data)
    except Exception as exc:
        print(f"  ! opportunities scan failed: {exc}")
        opportunities = []

    output = {
        "generated_at":    datetime.now(timezone.utc).isoformat(timespec='seconds'),
        "generator":       GENERATOR,
        "shared_context":  shared,
        "holdings":        holdings_results,
        "opportunities":   opportunities,
        "alert_summary":   summarise_alerts(holdings_results, opportunities),
    }

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    with OUTPUT_FILE.open("w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    print(f"  Wrote {OUTPUT_FILE.name}")

    archive_path = archive_report(OUTPUT_FILE)
    print(f"  Archived as {archive_path.name}")

    push_report_to_github(OUTPUT_FILE.parent)

    elapsed = (datetime.now() - started_at).total_seconds()
    print(f"  Summary: {output['alert_summary']}")
    print(f"  Elapsed: {elapsed:.1f}s")
    print("Done.")
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except Exception:
        traceback.print_exc()
        sys.exit(2)
