"""
generate_template_docx.py  —  one-shot helper

Creates a sample Holdings.docx in the private folder. Run this ONCE to seed a
starter Word document that the user can open, edit (replace example rows with
real holdings), and save. The lens's importer reads this file each run.

This file is NOT part of the lens runtime — it's a setup utility.

Run:
    py generate_template_docx.py
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PRIVATE_DIR = Path(os.environ.get(
    "HOLDINGS_PRIVATE_DIR",
    Path.home() / "OneDrive" / "Documents" / "Private Investments"
))
OUTPUT_PATH = PRIVATE_DIR / "Holdings.docx"


HOLDINGS_HEADERS = ["Name", "Type", "Value GBP", "Purchased", "Standing Alerts", "Notes"]
HOLDINGS_EXAMPLES = [
    ["Fundsmith Equity Class I Accumulation", "OEIC", "12000", "2024-03-15", "", "Long-term core holding"],
    ["Polar Capital Global Technology Class I", "OEIC", "5000", "2024-08-10", "High volatility", "Tech-heavy, accept the swings"],
    ["Scottish Mortgage Investment Trust", "InvestmentTrust", "3500", "2024-09-22", "", ""],
]

WATCHLIST_HEADERS = ["Name", "Type", "Rationale"]
WATCHLIST_EXAMPLES = [
    ["Artemis Global Income Class I", "OEIC", "Top of HL income-fund rankings; rotate in if growth fades"],
    ["Vanguard Global Small-Cap Index", "OEIC", "Recently added to HL Wealth Shortlist; consider for diversification"],
]


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    for r, row_data in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    return table


def main():
    PRIVATE_DIR.mkdir(parents=True, exist_ok=True)

    if OUTPUT_PATH.exists():
        print(f"REFUSING to overwrite existing file: {OUTPUT_PATH}")
        print("Delete or rename it first if you want a fresh template.")
        return 1

    doc = Document()

    title = doc.add_heading("My Investment Holdings", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "This is your private holdings record. The research lens reads this file each run "
        "and produces a fresh report. Edit the tables below — replace the example rows with "
        "your real holdings. Save the file when done. Never share this document."
    ).italic = True

    doc.add_heading("Current Holdings", level=1)
    doc.add_paragraph(
        "One row per holding. Required column: Name. Other columns are optional but "
        "useful — Value GBP enables portfolio-weight context, Standing Alerts flow into "
        "the lens's signal narrative."
    )
    add_table(doc, HOLDINGS_HEADERS, HOLDINGS_EXAMPLES)

    doc.add_paragraph()  # spacer

    doc.add_heading("Watchlist", level=1)
    doc.add_paragraph(
        "Funds you are watching but haven't bought yet. The lens uses Rationale to focus "
        "its research — e.g. 'rotate in if growth fades' tells the lens what would trigger your buy."
    )
    add_table(doc, WATCHLIST_HEADERS, WATCHLIST_EXAMPLES)

    doc.add_paragraph()
    notes = doc.add_paragraph()
    notes_run = notes.add_run(
        "Notes: Type values can be OEIC, ETF, InvestmentTrust, Equity, or Other. "
        "Dates in ISO format (YYYY-MM-DD). Leave a cell blank if the field doesn't apply. "
        "Do not add or remove columns — the importer expects the exact header names above."
    )
    notes_run.italic = True
    notes_run.font.size = Pt(9)
    notes_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT_PATH)
    print(f"Created template: {OUTPUT_PATH}")
    print("Open it in Word, replace the example rows with your real holdings, save.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
